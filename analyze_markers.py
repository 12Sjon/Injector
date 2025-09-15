"""
Analyze bullet markers in the enhanced resume preview files
Shows exactly what markers are used for new points
"""

from docx import Document
import sys

def analyze_bullet_markers(file_path, resume_name):
    print(f'\n{"="*70}')
    print(f'📄 ANALYZING BULLET MARKERS - {resume_name}')
    print(f'{"="*70}')
    
    try:
        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Find all bullet patterns
        bullet_patterns = {}
        bullet_examples = {}
        
        for i, para in enumerate(paragraphs):
            if para:
                first_char = para[0]
                if first_char in ['-', '•', '*', '◦', '▪', '‣', '⁃', '○', '▫']:
                    if first_char not in bullet_patterns:
                        bullet_patterns[first_char] = 0
                        bullet_examples[first_char] = []
                    bullet_patterns[first_char] += 1
                    if len(bullet_examples[first_char]) < 3:
                        bullet_examples[first_char].append(para[:80] + '...' if len(para) > 80 else para)
        
        print(f'📊 BULLET MARKER ANALYSIS:')
        if bullet_patterns:
            for marker, count in sorted(bullet_patterns.items(), key=lambda x: x[1], reverse=True):
                print(f'   "{marker}" : {count} occurrences')
        else:
            print('   No bullet markers found')
            
        print(f'\n🔸 BULLET EXAMPLES BY MARKER:')
        for marker, examples in bullet_examples.items():
            print(f'\n   📌 Marker: "{marker}" ({bullet_patterns[marker]} total uses)')
            for j, example in enumerate(examples, 1):
                print(f'      {j}. {example}')
        
        # Look for tech-specific content that starts with bullets
        tech_keywords = ['Java', 'Spring', 'AWS', 'React', 'microservices', 'REST', 'API', 'Lambda', 'EC2', 'Redux', 'enterprise', 'developed', 'implemented', 'created', 'configured', 'utilized']
        tech_bullets = []
        
        for para in paragraphs:
            if bullet_patterns:  # Only if we found bullet markers
                if any(keyword.lower() in para.lower() for keyword in tech_keywords) and para.startswith(tuple(bullet_patterns.keys())):
                    tech_bullets.append(para)
        
        print(f'\n💻 NEWLY ADDED TECH BULLETS: {len(tech_bullets)}')
        if tech_bullets:
            print('   (These are likely the new points added by RSInjector)')
            for i, bullet in enumerate(tech_bullets[:10], 1):  # Show first 10
                marker = bullet[0] if bullet else '?'
                content = bullet[:120] + '...' if len(bullet) > 120 else bullet
                print(f'   {i:2d}. 🆕 [{marker}] {content}')
            
            if len(tech_bullets) > 10:
                print(f'   ... and {len(tech_bullets) - 10} more tech bullets')
        else:
            print('   No tech-enhanced bullets detected with standard markers')
        
        # Show format pattern analysis
        print(f'\n🎯 MARKER FORMAT PATTERN:')
        if bullet_patterns:
            most_common = max(bullet_patterns.items(), key=lambda x: x[1])
            print(f'   Most used marker: "{most_common[0]}" ({most_common[1]} times)')
            print(f'   Format pattern: "{most_common[0]} [space] content"')
            print(f'   Example format: "{most_common[0]} • Developed enterprise-grade applications..."')
        
        return bullet_patterns, tech_bullets
            
    except Exception as e:
        print(f'❌ Error analyzing {resume_name}: {e}')
        return {}, []

def compare_original_vs_enhanced():
    """Compare original files with enhanced versions"""
    print(f'\n{"="*70}')
    print('🔍 COMPARISON: ORIGINAL vs ENHANCED')
    print(f'{"="*70}')
    
    # Files to compare
    files = [
        ('C:\\Users\\HP\\Downloads\\Resume Format 1.docx', 'preview_Resume_Format_1.docx', 'Resume Format 1'),
        ('C:\\Users\\HP\\Downloads\\Resume Format 3.DOCX', 'preview_Resume_Format_3.docx', 'Resume Format 3')
    ]
    
    for original_path, enhanced_path, name in files:
        print(f'\n📋 {name}:')
        
        try:
            # Analyze original
            orig_doc = Document(original_path)
            orig_paras = [p.text.strip() for p in orig_doc.paragraphs if p.text.strip()]
            orig_bullets = [p for p in orig_paras if p and p[0] in ['-', '•', '*', '◦', '▪', '‣', '⁃']]
            
            # Analyze enhanced
            enh_doc = Document(enhanced_path)
            enh_paras = [p.text.strip() for p in enh_doc.paragraphs if p.text.strip()]
            enh_bullets = [p for p in enh_paras if p and p[0] in ['-', '•', '*', '◦', '▪', '‣', '⁃']]
            
            print(f'   📊 Original bullets: {len(orig_bullets)}')
            print(f'   📊 Enhanced bullets: {len(enh_bullets)}')
            print(f'   🆕 New bullets added: {len(enh_bullets) - len(orig_bullets)}')
            
            # Show marker consistency
            if orig_bullets:
                orig_marker = orig_bullets[0][0]
                print(f'   📌 Original marker: "{orig_marker}"')
            else:
                print(f'   📌 Original marker: None (no bullets)')
                
            if enh_bullets:
                enh_marker = enh_bullets[0][0] if enh_bullets else None
                print(f'   📌 Enhanced marker: "{enh_marker}"')
                
                # Check consistency
                markers_used = set(b[0] for b in enh_bullets)
                if len(markers_used) == 1:
                    print(f'   ✅ Marker consistency: Perfect (all use "{list(markers_used)[0]}")')
                else:
                    print(f'   ⚠️  Multiple markers used: {markers_used}')
            
        except Exception as e:
            print(f'   ❌ Error comparing {name}: {e}')

def main():
    """Main analysis function"""
    print('🔍 BULLET MARKER ANALYSIS FOR RSINJECTOR')
    print('='*70)
    
    # Analyze enhanced files
    patterns1, tech1 = analyze_bullet_markers('preview_Resume_Format_1.docx', 'Resume Format 1 (Enhanced)')
    patterns2, tech2 = analyze_bullet_markers('preview_Resume_Format_3.docx', 'Resume Format 3 (Enhanced)')
    
    # Compare with originals
    compare_original_vs_enhanced()
    
    # Final summary
    print(f'\n{"🎉 FINAL SUMMARY ":=^70}')
    print(f'📄 Resume Format 1: {len(tech1)} tech bullets added')
    print(f'📄 Resume Format 3: {len(tech2)} tech bullets added')
    print(f'📊 Total tech enhancements: {len(tech1) + len(tech2)}')
    
    # Show the exact format used
    if patterns1:
        marker1 = list(patterns1.keys())[0]
        print(f'\n🎯 FORMAT USED FOR NEW POINTS:')
        print(f'   Resume Format 1: "{marker1} [content]"')
    
    if patterns2:
        marker2 = list(patterns2.keys())[0]
        print(f'   Resume Format 3: "{marker2} [content]"')
    
    print(f'\n✅ Analysis completed!')

if __name__ == "__main__":
    main()